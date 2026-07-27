# -*- coding: utf-8 -*-
"""
Sinh docs/BAO-CAO.docx tu docs/BAO-CAO.md

Chay:  python docs/build-baocao-docx.py
Can:   pip install python-docx

Diem dac biet: bang dac ta Use Case trong markdown duoc viet PHANG
(bang 4 dong + khoi "Kich ban chinh" + khoi "Ngoai le") vi markdown khong
long duoc bang trong o. Khi sinh docx, script GHEP lai thanh DUNG bang 6 dong
theo mau cua thay, voi bang du lieu LONG trong o "Kich ban chinh".
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, 'docs', 'BAO-CAO.md')
OUT = os.path.join(ROOT, 'docs', 'BAO-CAO.docx')

FONT = 'Times New Roman'
SIZE = 13
MONO = 'Consolas'

DAC_TA_KEYS = ('use case', 'actor', 'tiền điều kiện', 'hậu điều kiện')

# ---------------------------------------------------------------- tien ich

def strip_md(s):
    """Bo cu phap markdown con lai trong mot doan text tho."""
    s = s.replace('<br>', '\n').replace('<br/>', '\n')
    s = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', s)
    s = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', s)
    return s.replace(r'\|', '|').strip()


def plain(s):
    """Bo het cu phap nhan manh, dung de so khop nhan o bang."""
    s = strip_md(s)
    s = re.sub(r'[*`_]+', '', s)
    return s.strip().rstrip(':').lower()


TOKEN = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')


def write_runs(par, text):
    """Ghi text vao paragraph, giu dinh dang **dam**, `code`, *nghieng*."""
    for part in TOKEN.split(strip_md(text)):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = par.add_run(part[1:-1]); r.font.name = MONO
            r.font.size = Pt(SIZE - 2)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), MONO)
        elif part.startswith('*') and part.endswith('*'):
            r = par.add_run(part[1:-1]); r.italic = True
        else:
            par.add_run(part)


def split_row(line):
    """Tach mot dong bang markdown thanh cac o, ton trong \\| da escape."""
    out, cur, i = [], '', 0
    while i < len(line):
        c = line[i]
        if c == '\\' and i + 1 < len(line) and line[i + 1] == '|':
            cur += '|'; i += 2; continue
        if c == '|':
            out.append(cur); cur = ''; i += 1; continue
        cur += c; i += 1
    out.append(cur)
    if out and not out[0].strip():
        out = out[1:]
    if out and not out[-1].strip():
        out = out[:-1]
    return [c.strip() for c in out]


IS_SEP = re.compile(r'^\s*\|[\s:\-|]+\|\s*$')


# ---------------------------------------------------------------- phan tich

def parse(md):
    """Chuyen markdown thanh danh sach block."""
    lines = md.splitlines()
    blocks, i = [], 0
    while i < len(lines):
        l = lines[i]
        s = l.strip()

        if s.startswith('```'):                       # khung phac thao / code
            i += 1; buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1
            blocks.append(('code', buf)); continue

        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            blocks.append(('h', (len(m.group(1)), m.group(2)))); i += 1; continue

        if s.startswith('|'):                         # bang
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not IS_SEP.match(lines[i]):
                    rows.append(split_row(lines[i].strip()))
                i += 1
            if rows:
                blocks.append(('table', rows))
            continue

        m = re.match(r'^!\[([^\]]*)\]\(<?([^>)]+)>?\)\s*$', s)
        if m:
            blocks.append(('img', (m.group(1), m.group(2)))); i += 1; continue

        if s.startswith('>'):
            buf = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                buf.append(lines[i].strip().lstrip('>').strip()); i += 1
            blocks.append(('quote', ' '.join(x for x in buf if x))); continue

        if re.match(r'^([-*+]|\d+\.)\s+', s):
            blocks.append(('li', s)); i += 1; continue

        if s == '---':
            blocks.append(('hr', None)); i += 1; continue

        if s:
            blocks.append(('p', s))
        i += 1
    return blocks


def is_dac_ta(block):
    """Bang 4 dong dau cua dac ta Use Case?"""
    if block[0] != 'table':
        return False
    keys = [plain(r[0]) for r in block[1] if r]
    return all(any(k == want for k in keys) for want in DAC_TA_KEYS)


# ---------------------------------------------------------------- ket xuat

def add_par(doc, text='', style=None):
    p = doc.add_paragraph(style=style)
    if text:
        write_runs(p, text)
    return p


def add_code(container, lines):
    """Khung phac thao: chu don cach, giu nguyen khoang trang."""
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2); pf.space_after = Pt(2); pf.line_spacing = 1.0
    r = p.add_run('\n'.join(lines))
    r.font.name = MONO; r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), MONO)
    return p


def add_table(container, rows, nested=False):
    ncol = max(len(r) for r in rows)
    t = container.add_table(rows=len(rows), cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].text = ''
            txt = row[ci] if ci < len(row) else ''
            for k, seg in enumerate(strip_md(txt).split('\n')):
                p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                write_runs(p, seg)
                for r in p.runs:
                    r.font.size = Pt(SIZE - 2 if nested else SIZE - 1)
                    if ri == 0:
                        r.bold = True
    return t


def add_image(doc, alt, path):
    full = os.path.normpath(os.path.join(ROOT, 'docs', path))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(full):
        try:
            doc.paragraphs[-1]._p.getparent().remove(doc.paragraphs[-1]._p)
            doc.add_picture(full, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            write_runs(p, '[Loi chen anh: %s]' % e)
    else:
        r = p.add_run('[ CHUA CO HINH: %s ]' % os.path.basename(path))
        r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(alt); rc.italic = True; rc.font.size = Pt(SIZE - 2)


def emit_dac_ta(doc, head_rows, kich_ban, ngoai_le):
    """Ghep thanh DUNG bang 6 dong, bang du lieu LONG trong o Kich ban chinh."""
    t = doc.add_table(rows=6, cols=2)
    t.style = 'Table Grid'
    t.columns[0].width = Inches(1.4)
    t.columns[1].width = Inches(5.1)

    for ri, row in enumerate(head_rows):
        for ci in (0, 1):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].text = ''
            write_runs(cell.paragraphs[0], row[ci] if ci < len(row) else '')
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(SIZE - 1)

    # dong 5 - Kich ban chinh (co the chua bang long + khung phac thao)
    lbl = t.cell(4, 0); lbl.paragraphs[0].text = ''
    lbl.paragraphs[0].add_run('Kịch bản chính').bold = True
    body = t.cell(4, 1); body.paragraphs[0].text = ''
    first = True
    for kind, val in kich_ban:
        if kind == 'table':
            nt = add_table(body, val, nested=True)
            body.add_paragraph().paragraph_format.space_after = Pt(0)
        elif kind == 'code':
            add_code(body, val)
        else:
            p = body.paragraphs[0] if first else body.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            write_runs(p, val)
            for r in p.runs:
                r.font.size = Pt(SIZE - 1)
        first = False

    lbl2 = t.cell(5, 0); lbl2.paragraphs[0].text = ''
    lbl2.paragraphs[0].add_run('Ngoại lệ').bold = True
    nl = t.cell(5, 1); nl.paragraphs[0].text = ''
    first2 = True
    for kind, val in ngoai_le:
        if kind == 'table':
            add_table(nl, val, nested=True)
            nl.add_paragraph().paragraph_format.space_after = Pt(0)
        elif kind == 'code':
            add_code(nl, val)
        else:
            p = nl.paragraphs[0] if first2 else nl.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            write_runs(p, val)
            for r in p.runs:
                r.font.size = Pt(SIZE - 1)
        first2 = False
    doc.add_paragraph()


def build():
    if not os.path.exists(SRC):
        sys.exit('Khong thay %s' % SRC)
    blocks = parse(open(SRC, encoding='utf8').read())

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = FONT; st.font.size = Pt(SIZE)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.paragraph_format.line_spacing = 1.5
    st.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.left_margin = Inches(1.2); s.right_margin = Inches(0.8)
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)

    i, n_img, n_missing, n_dacta = 0, 0, 0, 0
    while i < len(blocks):
        kind, val = blocks[i]

        # --- gop dac ta Use Case thanh 1 bang 6 dong ---
        if is_dac_ta(blocks[i]):
            head = [r for r in val if plain(r[0]) in DAC_TA_KEYS]
            j = i + 1
            kich_ban, ngoai_le, doc_notes, mode = [], [], [], None
            while j < len(blocks):
                k2, v2 = blocks[j]
                label = plain(v2) if isinstance(v2, str) else ''
                if k2 in ('p', 'li') and label.startswith('kịch bản chính'):
                    mode = 'kb'; j += 1; continue
                if k2 in ('p', 'li') and label.startswith('ngoại lệ'):
                    mode = 'nl'; j += 1; continue
                if k2 == 'h' or (k2 == 'table' and is_dac_ta(blocks[j])):
                    break
                if mode is None and k2 == 'quote':
                    # ghi chu quy uoc ky hieu co the nam giua bang va Kich ban
                    doc_notes.append(v2); j += 1; continue
                if mode == 'kb':
                    kich_ban.append((k2, v2))
                elif mode == 'nl':
                    if k2 == 'quote':
                        break          # dong ghi chu anh xa -> ket thuc bang
                    ngoai_le.append((k2, v2))
                else:
                    break
                j += 1
            if mode is not None and len(head) == 4:
                for note in doc_notes:
                    q = add_par(doc, note)
                    q.paragraph_format.left_indent = Inches(0.3)
                    for r in q.runs:
                        r.italic = True; r.font.size = Pt(SIZE - 1)
                emit_dac_ta(doc, head, kich_ban, ngoai_le)
                n_dacta += 1
                i = j
                continue

        if kind == 'h':
            lvl, text = val
            if lvl == 1:
                if len(doc.paragraphs) > 1:
                    doc.add_page_break()
                p = doc.add_paragraph(style='Heading 1')
            else:
                p = doc.add_paragraph(style='Heading %d' % min(lvl, 4))
            write_runs(p, text)
            for r in p.runs:
                r.font.name = FONT
                r.font.color.rgb = RGBColor(0, 0, 0)
        elif kind == 'table':
            add_table(doc, val); doc.add_paragraph()
        elif kind == 'code':
            add_code(doc, val)
        elif kind == 'img':
            alt, path = val
            full = os.path.normpath(os.path.join(ROOT, 'docs', path))
            if not os.path.exists(full):
                n_missing += 1
            else:
                n_img += 1
            add_image(doc, alt, path)
        elif kind == 'quote':
            p = add_par(doc, val)
            p.paragraph_format.left_indent = Inches(0.3)
            for r in p.runs:
                r.italic = True; r.font.size = Pt(SIZE - 1)
        elif kind == 'li':
            p = add_par(doc, re.sub(r'^([-*+]|\d+\.)\s+', '', val))
            p.paragraph_format.left_indent = Inches(0.3)
        elif kind == 'p':
            add_par(doc, val)
        i += 1

    doc.save(OUT)
    print('Da sinh: %s' % OUT)
    print('  - bang dac ta Use Case ghep 6 dong (co bang long): %d' % n_dacta)
    print('  - anh chen duoc: %d' % n_img)
    print('  - anh CHUA CO (danh dau do trong file): %d' % n_missing)


if __name__ == '__main__':
    build()
